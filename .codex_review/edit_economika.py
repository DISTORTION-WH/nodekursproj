from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph


BASE = Path.cwd()
SRC = next(p for p in (BASE / "economika").glob("*.docx") if "02_27" in p.name)
OUT = SRC.with_name(SRC.stem + "_исправлено.docx")


def clear_paragraph(paragraph: Paragraph) -> None:
    p_pr = paragraph._p.pPr
    for child in list(paragraph._p):
        if child is not p_pr:
            paragraph._p.remove(child)


def set_paragraph(paragraph: Paragraph, text: str, *, bold: bool = True) -> None:
    alignment = paragraph.alignment
    style = paragraph.style
    clear_paragraph(paragraph)
    paragraph.style = style
    paragraph.alignment = alignment
    run = paragraph.add_run(text)
    run.bold = bold


def set_cell(cell, text: str, *, bold: bool = True) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold


def bold_cell(cell) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True


def find_paragraph(doc: Document, needle: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if needle in paragraph.text:
            return paragraph
    raise ValueError(f"Paragraph not found: {needle}")


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def delete_matching_paragraphs(doc: Document, needles: list[str]) -> None:
    for needle in needles:
        for paragraph in list(doc.paragraphs):
            if needle in paragraph.text:
                delete_paragraph(paragraph)
                break


def delete_table_rows_by_prefix(table: Table, prefixes: list[str]) -> None:
    for prefix in prefixes:
        for row in list(table.rows):
            if row.cells and row.cells[0].text.strip().startswith(prefix):
                table._tbl.remove(row._tr)
                break


def find_row(table: Table, prefix: str):
    for row in table.rows:
        if row.cells and row.cells[0].text.strip().startswith(prefix):
            return row
    raise ValueError(f"Row not found: {prefix}")


def insert_paragraph_after(paragraph: Paragraph, text: str = "", *, bold: bool = True) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if text:
        run = new_paragraph.add_run(text)
        run.bold = bold
    return new_paragraph


def insert_table_after(paragraph: Paragraph, rows: list[list[str]]) -> Table:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for row_idx, row_values in enumerate(rows):
        for col_idx, value in enumerate(row_values):
            cell = table.cell(row_idx, col_idx)
            cell.text = ""
            run = cell.paragraphs[0].add_run(value)
            run.bold = True
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
    paragraph._p.addnext(table._tbl)
    return table


def insert_paragraph_after_table(table: Table, text: str, *, bold: bool = True) -> Paragraph:
    new_p = OxmlElement("w:p")
    table._tbl.addnext(new_p)
    new_paragraph = Paragraph(new_p, table._parent)
    run = new_paragraph.add_run(text)
    run.bold = bold
    return new_paragraph


def set_table_widths(table: Table, widths_cm: list[float]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(width)


def bold_row(row) -> None:
    for cell in row.cells:
        bold_cell(cell)


doc = Document(SRC)

# 6.1: tighten the monetization sentence so it is explicitly one short characteristic.
set_paragraph(
    find_paragraph(doc, "В качестве модели монетизации принята freemium-модель"),
    "В качестве модели монетизации принята freemium-модель: базовые функции остаются бесплатными, а расширенные возможности, включая перевод сообщений, субтитры и персонализацию профиля, предоставляются по подписке.",
)

# Table 6.1: remove machine-time inputs and cap overhead at 30%.
initial = doc.tables[0]
delete_table_rows_by_prefix(
    initial,
    [
        "Стоимость одного машино-часа",
        "Норматив расхода машинного времени",
    ],
)
row_overhead = find_row(initial, "Норматив общепроизводственных")
set_cell(row_overhead.cells[3], "30")
bold_row(row_overhead)

set_paragraph(
    find_paragraph(doc, "Исходные данные подобраны с учетом методики"),
    "Исходные данные подобраны с учетом фактических условий дипломного проекта и нормативов, применяемых для учебного экономического обоснования программного средства. Норматив общепроизводственных и общехозяйственных расходов принят равным 30%, что соответствует допустимому диапазону 15-30%.",
)

# Labor and salary clarifications.
set_paragraph(
    find_paragraph(doc, "Это значение используется при расчете заработной платы и машинного времени."),
    "Общая трудоемкость разработки составляет 95 рабочих дней. Это значение используется при расчете заработной платы и последующих затрат.",
)
set_paragraph(
    find_paragraph(doc, "Среднемесячные значения заработной платы приняты экспертно"),
    "Основная заработная плата определяется по ролям, участвовавшим в создании программного средства. Среднемесячные значения заработной платы приняты по результатам исследования уровня оплаты труда для используемых технологий React, TypeScript, Node.js и PostgreSQL: разработчик - 1 800 руб., администратор баз данных - 2 000 руб., дизайнер - 1 700 руб., тестировщик - 1 500 руб. Количество рабочих дней в месяце принято равным 21.",
)

# Machine time: do not calculate without real purchase/payment.
set_paragraph(
    find_paragraph(doc, "Расходы на машинное время определяются по методике"),
    "Расходы на оплату машинного времени в расчете не начисляются, так как при разработке Lume Chat машинное время сторонних вычислительных мощностей не покупалось и фактически не оплачивалось. По данной статье затраты составляют 0,00 руб.",
)
delete_matching_paragraphs(
    doc,
    [
        "Смв = См.ч",
        "Смв = 0,05",
    ],
)
set_paragraph(
    find_paragraph(doc, "Спз = Соз ∙ Нп.з. / 100. (6.8)"),
    "Спз = Соз ∙ Нп.з. / 100. (6.7)",
)

# Overhead and dependent calculations.
set_paragraph(
    find_paragraph(doc, "В расчете принят минимальный норматив 50%"),
    "Общепроизводственные и общехозяйственные расходы относятся на конкретное программное средство по нормативу от основной заработной платы исполнителей. В расчете принят норматив 30%, не превышающий допустимый диапазон 15-30%.",
)
set_paragraph(
    find_paragraph(doc, "Собп,обх = Соз ∙ Нобп,обх / 100. (6.9)"),
    "Собп,обх = Соз ∙ Нобп,обх / 100. (6.8)",
)
set_paragraph(
    find_paragraph(doc, "Собп,обх = 8 028,57 ∙ 50 / 100"),
    "Собп,обх = 8 028,57 ∙ 30 / 100 = 2 408,57 руб.",
)
set_paragraph(
    find_paragraph(doc, "Ср = Соз + Сдз + Сфсзн + Сбгс + См + Ссопу + Смв"),
    "Ср = Соз + Сдз + Сфсзн + Сбгс + См + Ссопу + Спз + Собп,обх. (6.9)",
)
set_paragraph(
    find_paragraph(doc, "Ср = 8 028,57 + 1 204,29 + 3 139,17 + 55,40"),
    "Ср = 8 028,57 + 1 204,29 + 3 139,17 + 55,40 + 55,20 + 0,00 + 1 605,71 + 2 408,57 = 16 496,91 руб.",
)
set_paragraph(
    find_paragraph(doc, "Ср.с.а = Ср ∙ Нр.с.а. / 100.  (6.11)"),
    "Ср.с.а = Ср ∙ Нр.с.а. / 100.  (6.10)",
)
set_paragraph(
    find_paragraph(doc, "Ср.с.а = 18 172,11 ∙ 10 / 100"),
    "Ср.с.а = 16 496,91 ∙ 10 / 100 = 1 649,69 руб.",
)
set_paragraph(
    find_paragraph(doc, "Сп = Ср + Ср.с.а. (6.12)"),
    "Сп = Ср + Ср.с.а. (6.11)",
)
set_paragraph(
    find_paragraph(doc, "Сп = 18 172,11 + 1 817,21"),
    "Сп = 16 496,91 + 1 649,69 = 18 146,60 руб.",
)

# 6.16: add required economic-effect and analog analysis before pricing.
heading_616 = find_paragraph(doc, "6.16 Определение цены и оценка эффективности")
p1 = insert_paragraph_after(
    heading_616,
    "Разрабатываемое программное средство предназначено для потребления другими пользователями: учебными группами, небольшими командами и онлайн-сообществами. Экономический эффект разработчика формируется за счет реализации программного средства и регулярных подписочных платежей, а пользовательский эффект выражается в сокращении затрат времени на коммуникацию за счет объединения чатов, файлов, звонков, модерации, субтитров и перевода сообщений в одном веб-сервисе.",
)
p2 = insert_paragraph_after(
    p1,
    "Для выбора способа монетизации проанализированы программные продукты-аналоги Discord, Slack и Microsoft Teams. Указанные сервисы применяют модель бесплатного базового доступа с платными расширенными возможностями или подписками, поэтому для Lume Chat выбрана freemium-модель с платной подпиской на дополнительные функции.",
)
p_caption = insert_paragraph_after(p2, "Таблица 6.6 – Аналоги и модель монетизации Lume Chat")
p_caption.alignment = WD_ALIGN_PARAGRAPH.LEFT
analogs_table = insert_table_after(
    p_caption,
    [
        ["Продукт-аналог", "Разработчик и ссылка", "Модель монетизации", "Сравнение с Lume Chat"],
        ["Discord", "Discord Inc., discord.com", "Базовый доступ бесплатно, расширенные возможности по подписке", "Сходен по чатам, сообществам и звонкам; Lume Chat ориентирован на учебные группы и встроенные инструменты перевода"],
        ["Slack", "Salesforce, slack.com", "Бесплатный тариф и платные рабочие пространства", "Сходен по командным каналам, файлам и интеграциям; Lume Chat делает акцент на видеосвязи и модерации в учебных сообществах"],
        ["Microsoft Teams", "Microsoft, microsoft.com/microsoft-teams", "Бесплатный доступ и платные планы Microsoft 365", "Сходен по групповым коммуникациям и видеозвонкам; Lume Chat является более узким веб-сервисом с персонализацией и субтитрами"],
    ],
)
set_table_widths(analogs_table, [3.0, 4.0, 4.0, 5.0])
p3 = insert_paragraph_after_table(
    analogs_table,
    "Ключевыми характеристиками для сравнения выбраны групповые чаты, аудио- и видеозвонки, обмен файлами, модерация, браузерный доступ, субтитры, перевод сообщений и персонализация профиля. Эти характеристики определяют полезность коммуникационного сервиса для выбранной аудитории.",
)

# Keep the existing pricing-by-market-estimate block, but update dependent numbers.
set_paragraph(
    find_paragraph(doc, "Так как полная себестоимость разработки и первичного сопровождения ПС составляет"),
    "Так как полная себестоимость разработки и первичного сопровождения ПС составляет 18 146,60 руб., а средняя рыночная цена аналогов равна 32 091,75 руб., цену реализации целесообразно определить затратным методом с учетом заданного уровня рентабельности 30%. Расчетная цена с НДС находится в диапазоне рыночных оценок и ниже средней цены аналогов.",
)
set_paragraph(
    find_paragraph(doc, "ППС = Сп · Урент / 100"),
    "ППС = Сп · Урент / 100 = 18 146,60 · 30 / 100 = 5 443,98 руб.",
)
set_paragraph(
    find_paragraph(doc, "Цбез НДС = Сп + ППС"),
    "Цбез НДС = Сп + ППС = 18 146,60 + 5 443,98 = 23 590,58 руб.",
)
set_paragraph(
    find_paragraph(doc, "НДС = 25 986,12"),
    "НДС = 23 590,58 · 20 / 100 = 4 718,12 руб.;",
)
set_paragraph(
    find_paragraph(doc, "Цс НДС = 25 986,12"),
    "Цс НДС = 23 590,58 + 4 718,12 = 28 308,70 руб.",
)
set_paragraph(
    find_paragraph(doc, "РПС = (5 996,80 / 19 989,32)"),
    "РПС = (5 443,98 / 18 146,60) · 100 = 30,00%.",
)
set_paragraph(
    find_paragraph(doc, "ЧП = 5 996,80"),
    "ЧП = 5 443,98 · (1 - 20 / 100) = 4 355,18 руб.",
)
set_paragraph(
    find_paragraph(doc, "После учета годовых эксплуатационных расходов 1 560,00 руб."),
    "После учета годовых эксплуатационных расходов 1 560,00 руб. годовой результат до налога составляет 5 640,00 руб.; срок покрытия полной себестоимости за счет подписок - 18 146,60 / 5 640,00 = 3,22 года. Поэтому подписочная модель рассматривается как последующий источник регулярной выручки, а основная экономическая целесообразность подтверждается расчетной ценой реализации ПС.",
)

# Final table caption changes because a new table was inserted into 6.16.
set_paragraph(
    find_paragraph(doc, "Таблица 6.6 – Результаты расчетов"),
    "Таблица 6.7 – Результаты расчетов",
)

summary = doc.tables[-1]
updates = {
    "Общепроизводственные и общехозяйственные расходы": "2 408,57",
    "Сумма расходов на разработку ПС": "16 496,91",
    "Расходы на сопровождение и адаптацию": "1 649,69",
    "Полная себестоимость ПС": "18 146,60",
    "Принятая цена реализации с НДС": "28 308,70",
    "Цена программного средства без НДС": "23 590,58",
    "Прибыль от реализации ПС": "5 443,98",
    "Чистая прибыль": "4 355,18",
    "Срок покрытия полной себестоимости": "3,22",
}
for row in summary.rows:
    label = row.cells[0].text.strip()
    if label.startswith("Материалы, оборудование"):
        set_cell(row.cells[0], "Материалы, оборудование, платные услуги и машинное время, руб.")
        set_cell(row.cells[1], "55,20")
        bold_row(row)
    for prefix, value in updates.items():
        if label.startswith(prefix):
            set_cell(row.cells[1], value)
            bold_row(row)
            break

set_paragraph(
    find_paragraph(doc, "Полная себестоимость разработки и первичного сопровождения Lume Chat составляет"),
    "Полная себестоимость разработки и первичного сопровождения Lume Chat составляет 18 146,60 руб. По результатам анализа сайтов-агрегаторов минимальная рыночная оценка разработки аналогичного приложения составляет 26 526,50 руб., а средняя рыночная цена аналогов - 32 091,75 руб. В качестве цены реализации принимается расчетная цена 28 308,70 руб. с НДС, сформированная при уровне рентабельности 30%.",
)
set_paragraph(
    find_paragraph(doc, "Цена программного средства без НДС составляет"),
    "Цена программного средства без НДС составляет 23 590,58 руб., прибыль от реализации - 5 443,98 руб., рентабельность продукта - 30,00%, чистая прибыль - 4 355,18 руб. Годовая выручка по подписочной модели составляет 7 200,00 руб., а результат после эксплуатационных расходов - 5 640,00 руб. Следовательно, разработка Lume Chat является экономически целесообразной при реализации по расчетной цене разработчика, а подписочная модель может использоваться как дополнительный регулярный источник выручки.",
)

# Append a bold change log requested by the user.
doc.add_paragraph()
change_heading = doc.add_paragraph()
change_heading.add_run("Список внесенных изменений").bold = True

change_items = [
    "Удалено: из исходных данных и расчетов исключены стоимость одного машино-часа, норматив машинного времени и расчет машинного времени 69,48 руб., так как машинное время не покупалось и не оплачивалось.",
    "Добавлено: в раздел 6.16 добавлено описание экономического эффекта, целевой аудитории, выбранной freemium-модели и сравнение с программными продуктами-аналогами Discord, Slack и Microsoft Teams.",
    "Изменено: норматив общепроизводственных и общехозяйственных расходов уменьшен с 50% до 30%, что соответствует допустимому диапазону 15-30%.",
    "Изменено: пересчитаны общепроизводственные расходы, сумма расходов на разработку, расходы на сопровождение, полная себестоимость, цена без НДС, НДС, цена с НДС, прибыль, чистая прибыль и срок покрытия себестоимости за счет подписок.",
    "Изменено: раздел о заработной плате уточнен с привязкой к используемым технологиям React, TypeScript, Node.js и PostgreSQL.",
    "Изменено: итоговая таблица и выводы по разделу приведены в соответствие с пересчитанными показателями.",
]
for item in change_items:
    paragraph = doc.add_paragraph(style=None)
    run = paragraph.add_run(item)
    run.bold = True

doc.save(OUT)
print(OUT)
