from __future__ import annotations

from pathlib import Path
import re
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "zapiska" / "listingi_koda_prilozheniya.docx"


def clean_text(value: str) -> str:
    """Normalize document text to match the note formatting requirements."""
    return value.replace("ё", "е").replace("Ё", "Е").replace("\t", "    ")


def read_source(relative_path: str) -> str:
    return clean_text((ROOT / relative_path).read_text(encoding="utf-8"))


def line_start(text: str, index: int) -> int:
    return text.rfind("\n", 0, index) + 1


def scan_matching_brace(text: str, open_index: int) -> int:
    depth = 0
    i = open_index
    quote: str | None = None
    escape = False
    line_comment = False
    block_comment = False

    while i < len(text):
        ch = text[i]
        nxt = text[i + 1] if i + 1 < len(text) else ""

        if line_comment:
            if ch == "\n":
                line_comment = False
            i += 1
            continue

        if block_comment:
            if ch == "*" and nxt == "/":
                block_comment = False
                i += 2
            else:
                i += 1
            continue

        if quote is not None:
            if escape:
                escape = False
            elif ch == "\\":
                escape = True
            elif ch == quote:
                quote = None
            i += 1
            continue

        if ch == "/" and nxt == "/":
            line_comment = True
            i += 2
            continue
        if ch == "/" and nxt == "*":
            block_comment = True
            i += 2
            continue
        if ch in ("'", '"', "`"):
            quote = ch
            i += 1
            continue

        if ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                return i
        i += 1

    raise ValueError("Matching brace was not found")


def find_signature(text: str, name: str, extra_patterns: Iterable[str] = ()) -> int:
    patterns = [
        rf"(?m)^\s*(?:(?:public|private|protected)\s+)?(?:async\s+)?{re.escape(name)}\s*\(",
        rf"(?m)^\s*(?:async\s+)?function\s+{re.escape(name)}\s*\(",
        rf"(?m)^\s*const\s+{re.escape(name)}\s*=",
        rf"(?m)^\s*export\s+(?:async\s+)?function\s+{re.escape(name)}\s*\(",
        rf"(?m)^\s*export\s+const\s+{re.escape(name)}\s*=",
        *extra_patterns,
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            return match.start()
    raise ValueError(f"Signature not found: {name}")


def extract_named_block(relative_path: str, name: str, extra_patterns: Iterable[str] = ()) -> str:
    text = read_source(relative_path)
    start = find_signature(text, name, extra_patterns)
    start = line_start(text, start)
    open_index = text.find("{", start)
    if open_index == -1:
        raise ValueError(f"Opening brace not found for {name}")
    close_index = scan_matching_brace(text, open_index)

    end = close_index + 1
    # Class methods end at the brace. Assignments/useCallback blocks need the
    # following semicolon or closing call to remain readable as code fragments.
    tail = text[end : end + 500]
    if re.match(r"\s*[,)]", tail) or re.match(r"\s*;", tail):
        semicolon = text.find(";", end)
        if semicolon != -1:
            end = semicolon + 1
    elif re.match(r"\s*=>", text[start:open_index]):
        semicolon = text.find(";", end)
        if semicolon != -1:
            end = semicolon + 1

    return text[start:end].lstrip("\n").rstrip()


def extract_socket_event(relative_path: str, event_name: str) -> str:
    text = read_source(relative_path)
    token = f'socket.on("{event_name}"'
    start = text.find(token)
    if start == -1:
        raise ValueError(f"Socket event not found: {event_name}")
    start = line_start(text, start)
    arrow_index = text.find("=>", start)
    if arrow_index == -1:
        raise ValueError(f"Socket event callback not found: {event_name}")
    open_index = text.find("{", arrow_index)
    close_index = scan_matching_brace(text, open_index)
    semicolon = text.find(";", close_index)
    end = semicolon + 1 if semicolon != -1 else close_index + 1
    return text[start:end].lstrip("\n").rstrip()


def extract_call(relative_path: str, call_start: str) -> str:
    text = read_source(relative_path)
    start = text.find(call_start)
    if start == -1:
        raise ValueError(f"Call not found: {call_start}")
    start = line_start(text, start)
    open_index = text.find("{", start)
    close_index = scan_matching_brace(text, open_index)
    semicolon = text.find(";", close_index)
    end = semicolon + 1 if semicolon != -1 else close_index + 1
    return text[start:end].lstrip("\n").rstrip()


def full_file(relative_path: str) -> str:
    return read_source(relative_path).rstrip()


def set_cell_shading(paragraph, color: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    p_pr.append(shading)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def set_run_font(run, name: str, size: float, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    section.header_distance = Cm(1)

    header = section.header
    paragraph = header.paragraphs[0]
    add_page_number(paragraph)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)


def add_appendix_heading(document: Document, label: str, title: str) -> None:
    if len(document.paragraphs) > 1:
        document.add_section(WD_SECTION.NEW_PAGE)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(label)
    set_run_font(run, "Times New Roman", 14, True)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(12)
    run = paragraph.add_run(title)
    set_run_font(run, "Times New Roman", 14, False)


def add_listing_heading(document: Document, title: str, source: str | None = None) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(clean_text(title))
    set_run_font(run, "Times New Roman", 14, False)

    if source:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.keep_with_next = True
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(clean_text(f"Файл: {source}"))
        set_run_font(run, "Times New Roman", 14, False)


def add_code(document: Document, code: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.left_indent = Cm(0)
    paragraph.paragraph_format.right_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    run = paragraph.add_run(code if code else " ")
    set_run_font(run, "Courier New", 12, False)


def add_listing(document: Document, title: str, source: str, code: str) -> None:
    add_listing_heading(document, title, source)
    add_code(document, code)


def build() -> None:
    document = Document()
    configure_document(document)

    add_appendix_heading(document, "ПРИЛОЖЕНИЕ Ж", "Функции аутентификации и работы с пользователями")
    auth_file = "apps/api/Controllers/authController.ts"
    for name, title in [
        ("preRegister", "Функция preRegister"),
        ("login", "Функция login"),
    ]:
        add_listing(document, title, auth_file, extract_named_block(auth_file, name))

    user_file = "apps/api/Services/userService.ts"
    for name, title in [
        ("createUser", "Функция createUser"),
    ]:
        add_listing(document, title, user_file, extract_named_block(user_file, name))

    add_appendix_heading(document, "ПРИЛОЖЕНИЕ И", "Функции чатов, сообщений и групповых комнат")
    chat_service_file = "apps/api/Services/chatService.ts"
    for name, title in [
        ("getChatMessages", "Функция getChatMessages"),
        ("postMessage", "Функция postMessage"),
        ("forwardMessage", "Функция forwardMessage"),
    ]:
        add_listing(document, title, chat_service_file, extract_named_block(chat_service_file, name))

    chat_controller_file = "apps/api/Controllers/chatController.ts"
    for name, title in [
        ("postMessage", "Контроллер postMessage"),
    ]:
        add_listing(document, title, chat_controller_file, extract_named_block(chat_controller_file, name))

    add_appendix_heading(document, "ПРИЛОЖЕНИЕ К", "Функции администрирования, модерации и WebSocket-соединений")
    admin_file = "apps/api/Controllers/adminController.ts"
    for name, title in [
        ("getStats", "Функция getStats"),
    ]:
        add_listing(document, title, admin_file, extract_named_block(admin_file, name))

    moderator_file = "apps/api/Controllers/moderatorController.ts"
    for name, title in [
        ("banUser", "Функция banUser"),
    ]:
        add_listing(document, title, moderator_file, extract_named_block(moderator_file, name))

    index_file = "apps/api/index.ts"
    add_listing(document, "Middleware аутентификации Socket.io", index_file, extract_call(index_file, "io.use("))
    add_listing(document, "Функция checkCallRateLimit", index_file, extract_named_block(index_file, "checkCallRateLimit"))
    for event, title in [
        ("call_user", "Обработчик события call_user"),
    ]:
        add_listing(document, title, index_file, extract_socket_event(index_file, event))

    add_appendix_heading(document, "ПРИЛОЖЕНИЕ Л", "Функции видеозвонков и клиентской части приложения")
    mediasoup_file = "apps/api/Services/mediasoupService.ts"
    for name, title in [
        ("createWebRtcTransport", "Функция createWebRtcTransport"),
        ("produce", "Функция produce"),
        ("consume", "Функция consume"),
    ]:
        add_listing(document, title, mediasoup_file, extract_named_block(mediasoup_file, name))

    api_file = "apps/web/src/services/api.tsx"
    add_listing(document, "Обработчик ошибки 401 в Axios", api_file, extract_call(api_file, "async (error: AxiosError) => {"))

    chat_context_file = "apps/web/src/context/ChatContext.tsx"
    for name, title in [
        ("sendMessage", "Функция sendMessage"),
        ("loadMoreMessages", "Функция loadMoreMessages"),
    ]:
        add_listing(document, title, chat_context_file, extract_named_block(chat_context_file, name))

    call_context_file = "apps/web/src/context/CallContext.tsx"
    for name, title in [
        ("startCall", "Функция startCall"),
    ]:
        extra = (rf"(?m)^\s*const\s+{re.escape(name)}\s*=\s*useCallback",)
        add_listing(document, title, call_context_file, extract_named_block(call_context_file, name, extra))

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
