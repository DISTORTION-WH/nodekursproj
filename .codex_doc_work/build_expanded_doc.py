from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT_PATH = r".codex_doc_work\questions_expanded_ru_en.docx"


def set_cell_shading(paragraph, color):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    p_pr.append(shd)


def set_spacing(paragraph, before=0, after=6, line=1.08):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_label(doc, text, color="1F4E79"):
    p = doc.add_paragraph()
    set_spacing(p, before=6, after=2, line=1.0)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.size = Pt(10)
    return p


def add_text(doc, text, bold=False, italic=False, color=None, size=10.5, style=None):
    p = doc.add_paragraph(style=style)
    set_spacing(p, before=0, after=6, line=1.08)
    for idx, part in enumerate(text.split("\n")):
        if idx:
            p.add_run().add_break()
        run = p.add_run(part)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_spacing(p, before=0, after=3, line=1.06)
        run = p.add_run(item)
        run.font.size = Pt(10.2)


def add_answer(doc, blocks):
    for block in blocks:
        if isinstance(block, tuple) and block[0] == "bullets":
            add_bullets(doc, block[1])
        else:
            add_text(doc, block)


def add_qa(doc, index, q_ru, q_en, answer_ru, answer_en):
    p_num = doc.add_paragraph()
    set_spacing(p_num, before=10, after=2, line=1.0)
    r = p_num.add_run(f"{index}.")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string("1F4E79")

    add_label(doc, "Вопрос на русском")
    add_text(doc, q_ru, bold=True, size=11)

    add_label(doc, "English translation")
    add_text(doc, q_en, italic=True, color="4F4F4F", size=10.5)

    add_label(doc, "Ответ на русском")
    add_answer(doc, answer_ru)

    add_label(doc, "Answer in English")
    add_answer(doc, answer_en)


def add_section(doc, title):
    p = doc.add_paragraph()
    set_spacing(p, before=14, after=7, line=1.0)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string("1F4E79")
    set_cell_shading(p, "EAF2F8")


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(title, before=0, after=4, line=1.0)
    r = title.add_run("Вопросы и ответы: расширенная версия RU/EN")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor.from_string("17365D")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(subtitle, before=0, after=12, line=1.0)
    r = subtitle.add_run("Формулировки вопросов сохранены, ответы расширены и переведены на английский.")
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string("666666")

    add_section(doc, "JavaScript, TypeScript, NestJS, Testing")

    add_qa(
        doc,
        1,
        "1 Как работает scope chain? ",
        "1 How does the scope chain work?",
        [
            "Scope chain — это механизм поиска переменных в JavaScript. Когда код обращается к имени переменной, движок сначала ищет его в текущей области видимости. Если переменная там не найдена, поиск поднимается во внешний scope, затем еще выше, пока не дойдет до глобальной области. Если имя не найдено нигде, возникает ошибка ReferenceError.",
            "Важно, что scope в JavaScript лексический: он определяется местом, где функция была объявлена, а не местом, где она была вызвана. Поэтому вложенная функция «помнит» окружение, в котором была создана. Это лежит в основе замыканий: функция может использовать переменные внешней функции даже после того, как внешняя функция уже завершила выполнение.",
            "На практике это важно для callbacks, обработчиков событий, асинхронного кода и React hooks. Ошибки часто появляются, когда разработчик ожидает, что переменная будет искаться по месту вызова функции, но на самом деле она берется из окружения объявления.",
            ("bullets", [
                "let и const имеют блочную область видимости.",
                "var имеет функциональную область видимости и поднимается иначе.",
                "Scope chain отвечает за поиск переменных, но не определяет this.",
                "Замыкание возникает, когда функция сохраняет доступ к внешнему lexical scope."
            ]),
        ],
        [
            "The scope chain is JavaScript's mechanism for resolving variable names. When the engine sees a variable reference, it first looks in the current scope. If the variable is not found there, it moves to the outer scope, then further outward until it reaches the global scope. If the name cannot be resolved anywhere, JavaScript throws a ReferenceError.",
            "The key point is that JavaScript uses lexical scoping: the scope is determined by where a function is declared, not where it is called. This is why an inner function can remember variables from its outer function even after the outer function has already returned. That behavior is the basis of closures.",
            "In real applications this matters in callbacks, event handlers, asynchronous code, and React hooks. A common mistake is expecting a variable to be resolved from the call location, while JavaScript actually resolves it from the declaration environment.",
            ("bullets", [
                "let and const are block-scoped.",
                "var is function-scoped and behaves differently with hoisting.",
                "The scope chain resolves variables, but it does not determine this.",
                "A closure appears when a function keeps access to an outer lexical scope."
            ]),
        ],
    )

    add_qa(
        doc,
        2,
        "2 Как определяется значение this в JavaScript? ",
        "2 How is the value of this determined in JavaScript?",
        [
            "Значение this в JavaScript определяется способом вызова функции. Это главное отличие this от обычных переменных: переменные ищутся через scope chain, а this зависит от call site — места и формы вызова.",
            "Если функция вызвана как метод объекта, this обычно указывает на объект перед точкой: user.showName() дает this === user. Если метод сохранить в отдельную переменную и вызвать отдельно, связь с объектом теряется. В strict mode при обычном вызове функции this будет undefined, а без strict mode может стать глобальным объектом.",
            "this можно задать явно через call, apply или bind. При вызове через new создается новый объект, и this указывает на него. У стрелочных функций собственного this нет: они берут this из внешнего лексического окружения. Поэтому arrow functions удобны для callbacks, но часто плохи как методы объектов, если метод должен работать с this.",
            ("bullets", [
                "Вызов как метод: this — объект перед точкой.",
                "Обычный вызов: this зависит от strict mode.",
                "call/apply/bind задают this явно.",
                "new создает новый this-объект.",
                "Arrow function не имеет собственного this."
            ]),
        ],
        [
            "In JavaScript, this is determined by how a function is called. This is the main difference between this and normal variables: variables are resolved through the scope chain, while this is resolved from the call site.",
            "If a function is called as an object method, this usually refers to the object before the dot: user.showName() gives this === user. If the method is saved into a separate variable and called as a plain function, the object binding is lost. In strict mode, a plain function call has this === undefined; outside strict mode it may point to the global object.",
            "this can be set explicitly with call, apply, or bind. When a function is called with new, JavaScript creates a new object and binds this to it. Arrow functions do not have their own this; they capture this from the surrounding lexical environment. That makes them useful for callbacks, but often unsuitable as object methods when the method needs its own receiver.",
            ("bullets", [
                "Method call: this is the object before the dot.",
                "Plain function call: this depends on strict mode.",
                "call/apply/bind set this explicitly.",
                "new creates a new this object.",
                "Arrow functions do not define their own this."
            ]),
        ],
    )

    add_qa(
        doc,
        3,
        "3 Чем Map отличается от WeakMap? ",
        "3 How is Map different from WeakMap?",
        [
            "Map — это коллекция ключ-значение, где ключом может быть почти любое значение: объект, строка, число, функция и так далее. Map хранит сильные ссылки на ключи. Пока объект используется как ключ в Map, сборщик мусора не удалит этот объект, даже если других ссылок на него уже нет.",
            "WeakMap отличается тем, что ключами могут быть только объекты, а ссылки на ключи являются слабыми. Если объект больше нигде не используется, он может быть удален сборщиком мусора, даже если он был ключом в WeakMap. Это удобно для хранения метаданных, приватных данных или кэша, который не должен мешать очистке памяти.",
            "Из-за слабых ссылок WeakMap нельзя перебирать: у него нет size, keys, values, entries и for...of. JavaScript не может гарантировать стабильный список ключей, потому что часть объектов может быть удалена сборщиком мусора в любой момент.",
            ("bullets", [
                "Map подходит, когда нужно хранить и перебирать коллекцию.",
                "WeakMap подходит, когда данные привязаны к объекту и не должны создавать memory leak.",
                "Map может иметь примитивные ключи, WeakMap — только объектные.",
                "WeakMap полезен для DOM-метаданных, приватных данных и объектных кэшей."
            ]),
        ],
        [
            "Map is a key-value collection where a key can be almost any value: an object, string, number, function, and so on. Map keeps strong references to its keys. As long as an object is used as a key in a Map, the garbage collector cannot remove that object, even if there are no other references to it.",
            "WeakMap is different because its keys must be objects and those key references are weak. If an object is no longer referenced anywhere else, it can be garbage-collected even if it was used as a key in a WeakMap. This is useful for metadata, private object data, or caches that should not prevent memory cleanup.",
            "Because the references are weak, WeakMap cannot be iterated. It has no size, keys, values, entries, or for...of support. JavaScript cannot expose a stable list of keys because some of them may disappear after garbage collection.",
            ("bullets", [
                "Map is useful when the collection must be enumerable.",
                "WeakMap is useful when attached data must not cause memory leaks.",
                "Map can use primitive keys; WeakMap can only use object keys.",
                "WeakMap is useful for DOM metadata, private data, and object-level caches."
            ]),
        ],
    )

    add_qa(
        doc,
        4,
        "4 Почему type assertion через as может быть опасен?  ",
        "4 Why can type assertion through as be dangerous?",
        [
            "Type assertion через as не преобразует данные в runtime. Он только говорит TypeScript-компилятору: «считай это значение таким типом». Если разработчик ошибся, TypeScript перестанет предупреждать, но реальное значение останется прежним, и ошибка проявится уже в runtime.",
            "Самый опасный случай — данные с границ системы: API responses, Kafka messages, localStorage, process.env, OpenAI output, данные из БД после ручного mapping. Например, await response.json() as User не проверяет, что у объекта действительно есть email, id или role. Если backend вернул другую форму, код может упасть на user.email.toLowerCase().",
            "Особенно рискованны as any, двойное приведение as unknown as SomeType и non-null assertion через !. Они часто маскируют проблему вместо того, чтобы ее решить. Правильнее использовать runtime validation: Zod, class-validator, io-ts, type guards или ручную проверку структуры.",
            ("bullets", [
                "as — это подсказка компилятору, а не runtime validation.",
                "Для внешних данных лучше начинать с unknown и проверять форму.",
                "Type assertion допустим, когда разработчик действительно знает больше компилятора.",
                "Если as используется часто, это может быть сигналом слабой типовой модели."
            ]),
        ],
        [
            "A type assertion with as does not transform data at runtime. It only tells the TypeScript compiler: “treat this value as this type.” If the developer is wrong, TypeScript stops warning, but the actual runtime value remains unchanged and the bug appears later during execution.",
            "The most dangerous cases are system boundaries: API responses, Kafka messages, localStorage, process.env, OpenAI output, or database data after manual mapping. For example, await response.json() as User does not verify that the object really contains email, id, or role. If the backend returns a different shape, the code may crash on user.email.toLowerCase().",
            "The riskiest patterns are as any, double assertions like as unknown as SomeType, and the non-null assertion operator !. They often hide a problem instead of solving it. A safer approach is runtime validation with Zod, class-validator, io-ts, type guards, or explicit structural checks.",
            ("bullets", [
                "as is a compiler hint, not runtime validation.",
                "External data should usually start as unknown and then be validated.",
                "Type assertions are acceptable when the developer truly knows more than the compiler.",
                "Frequent use of as can indicate a weak type model."
            ]),
        ],
    )

    add_qa(
        doc,
        5,
        "5 Что такое cyclic dependency в NestJS? ",
        "5 What is a cyclic dependency in NestJS?",
        [
            "Cyclic dependency в NestJS — это круговая зависимость между providers или modules. Например, UserService зависит от AuthService, а AuthService одновременно зависит от UserService. На уровне модулей ситуация похожая: UserModule импортирует AuthModule, а AuthModule импортирует UserModule.",
            "Проблема в том, что NestJS должен построить dependency graph и создать providers в корректном порядке. При цикле порядок становится неоднозначным, могут появляться undefined providers, сложные ошибки DI и хрупкая архитектура. Даже если приложение запускается, такие зависимости ухудшают тестируемость и делают границы модулей неясными.",
            "NestJS предлагает forwardRef как workaround, но это не должно быть первым архитектурным решением. forwardRef помогает разорвать проблему инициализации, но не устраняет высокую связность. Более здоровые варианты: вынести общую логику в третий service, разделить ответственность, использовать domain events, ports/interfaces или пересмотреть модульные границы.",
            ("bullets", [
                "Cyclic dependency часто показывает, что два сервиса знают друг о друге слишком много.",
                "forwardRef допустим точечно, но не должен становиться стандартом.",
                "Лучшее решение обычно архитектурное: разделение ответственности или событийная коммуникация.",
                "На интервью важно объяснить не только workaround, но и причину проблемы."
            ]),
        ],
        [
            "A cyclic dependency in NestJS is a circular dependency between providers or modules. For example, UserService depends on AuthService while AuthService also depends on UserService. At module level the same pattern appears when UserModule imports AuthModule and AuthModule imports UserModule.",
            "The problem is that NestJS has to build a dependency graph and instantiate providers in a valid order. With a cycle, that order becomes ambiguous. This can lead to undefined providers, confusing DI errors, and fragile architecture. Even if the application starts, such dependencies reduce testability and make module boundaries unclear.",
            "NestJS provides forwardRef as a workaround, but it should not be the first architectural answer. forwardRef helps with initialization, but it does not fix high coupling. Healthier options include extracting shared logic into a third service, splitting responsibilities, using domain events, introducing ports/interfaces, or revisiting module boundaries.",
            ("bullets", [
                "A cyclic dependency often means that two services know too much about each other.",
                "forwardRef is acceptable in narrow cases, but should not become the default.",
                "The best fix is usually architectural: separate responsibilities or use event-based communication.",
                "In an interview, it is important to explain both the workaround and the underlying design issue."
            ]),
        ],
    )

    add_qa(
        doc,
        6,
        "6 Что означает принцип F.I.R.S.T. в тестировании? ",
        "6 What does the F.I.R.S.T. principle mean in testing?",
        [
            "F.I.R.S.T. — это набор свойств хороших unit-тестов: Fast, Independent, Repeatable, Self-validating, Timely. Идея принципа в том, что тесты должны давать быструю и надежную обратную связь, а не превращаться в нестабильный ритуал перед релизом.",
            "Fast означает, что тесты должны выполняться быстро. Если unit-тесты идут долго, разработчики запускают их реже. Independent означает, что тесты не зависят друг от друга и могут выполняться в любом порядке. Repeatable означает, что тест стабильно проходит локально, в CI и на другой машине, без зависимости от текущей даты, внешнего API или случайного состояния.",
            "Self-validating означает, что тест сам говорит pass/fail через assertions, а не требует ручной проверки логов. Timely означает, что тесты пишутся вовремя — до, во время или сразу после реализации, пока понятны edge cases и код еще можно легко спроектировать тестируемым.",
            ("bullets", [
                "Fast: быстрые тесты чаще запускают.",
                "Independent: тесты не зависят от порядка выполнения.",
                "Repeatable: результат стабилен в разных окружениях.",
                "Self-validating: тест сам проверяет результат.",
                "Timely: тесты пишутся вовремя, а не постфактум через месяц."
            ]),
        ],
        [
            "F.I.R.S.T. describes the qualities of good unit tests: Fast, Independent, Repeatable, Self-validating, and Timely. The idea is that tests should provide fast and reliable feedback instead of becoming an unstable release ritual.",
            "Fast means tests should run quickly. If unit tests are slow, developers run them less often. Independent means tests must not depend on each other and should pass in any order. Repeatable means the same test should pass locally, in CI, and on another machine without depending on the current date, external APIs, or shared random state.",
            "Self-validating means a test clearly reports pass or fail through assertions instead of requiring manual log inspection. Timely means tests should be written before, during, or immediately after implementation, while edge cases are still clear and the code can still be designed for testability.",
            ("bullets", [
                "Fast: fast tests are run more often.",
                "Independent: tests do not rely on execution order.",
                "Repeatable: results are stable across environments.",
                "Self-validating: the test checks the result itself.",
                "Timely: tests are written on time, not reconstructed months later."
            ]),
        ],
    )

    add_section(doc, "Frontend, React, Next.js")

    add_qa(
        doc,
        7,
        "Представь что страница весит один гигабайт, как со стороны пользователя ты бы оптимизировал это со стороны фронтенда",
        "Imagine a page weighs one gigabyte. From the user's perspective, how would you optimize it on the frontend?",
        [
            "Сначала нужно понять, что именно дает один гигабайт: JavaScript bundle, изображения, видео, шрифты, JSON responses, sourcemaps или сторонние скрипты. Для этого я бы использовал Chrome DevTools Network, Coverage, Performance, Lighthouse и bundle analyzer для Next.js. Без измерений легко оптимизировать не тот слой.",
            "В React и Next.js основной принцип — уменьшить initial load. Тяжелые компоненты нужно грузить через dynamic import или React.lazy: графики, редакторы, markdown preview, video player, карты, PDF viewer. В Next.js важно держать 'use client' как можно ниже, потому что client component попадает в браузерный bundle, а server component — нет.",
            "Для изображений использовать next/image, правильные sizes, WebP/AVIF, lazy loading и priority только для действительно critical image. Для видео — poster, preload='none' или preload='metadata', загрузку player по клику, HLS/DASH streaming вместо скачивания большого файла целиком. Для больших списков — pagination, infinite scroll и virtualization через react-window или аналог.",
            "В Next.js App Router полезны streaming и Suspense: пользователь быстро получает shell страницы, а тяжелые блоки догружаются позже. Для данных — server-side filtering, cursor pagination, кеширование fetch с revalidate, SWR/React Query на клиенте. Если есть тяжелая обработка данных в браузере, ее лучше вынести в Web Worker, чтобы не блокировать main thread.",
            ("bullets", [
                "Code splitting: грузить только код текущего экрана.",
                "Lazy loading: откладывать тяжелые компоненты, изображения и видео.",
                "Server Components: не отправлять на клиент то, что можно выполнить на сервере.",
                "Streaming/Suspense: показывать полезный UI до готовности всех данных.",
                "Virtualization: не рендерить тысячи элементов в DOM одновременно.",
                "Caching: повторно использовать assets и данные, где это безопасно."
            ]),
        ],
        [
            "The first step is to identify what actually creates the one-gigabyte payload: JavaScript bundles, images, video, fonts, JSON responses, sourcemaps, or third-party scripts. I would use Chrome DevTools Network, Coverage, Performance, Lighthouse, and a Next.js bundle analyzer. Without measurement, it is easy to optimize the wrong layer.",
            "In React and Next.js, the main principle is to reduce the initial load. Heavy components should be loaded through dynamic import or React.lazy: charts, rich text editors, markdown previews, video players, maps, and PDF viewers. In Next.js it is important to keep 'use client' as low as possible because client components are shipped to the browser, while server components are not.",
            "For images, I would use next/image, correct sizes, WebP/AVIF, lazy loading, and priority only for truly critical images. For video, I would use a poster, preload='none' or preload='metadata', load the player on demand, and use HLS/DASH streaming instead of downloading a large file upfront. For large lists, I would use pagination, infinite scroll, and virtualization with react-window or a similar library.",
            "In the Next.js App Router, streaming and Suspense are useful because the user can receive the page shell quickly while heavy blocks load later. For data, I would use server-side filtering, cursor pagination, fetch caching with revalidate, and SWR/React Query on the client. Heavy browser-side processing should be moved to a Web Worker so it does not block the main thread.",
            ("bullets", [
                "Code splitting: load only the code needed for the current screen.",
                "Lazy loading: defer heavy components, images, and video.",
                "Server Components: avoid sending server-only work to the client.",
                "Streaming/Suspense: show useful UI before all data is ready.",
                "Virtualization: do not render thousands of DOM nodes at once.",
                "Caching: reuse assets and data where it is safe."
            ]),
        ],
    )

    add_qa(
        doc,
        8,
        "Задачка с библиотекой crypto.  Представь, ты работаешь с crypto module, запустил подряд 4 команды шифрования пароля, за какое время выполнится каждая из этих команд и будет ли значительно отличаться это время? ",
        "A task with the crypto library. Imagine you are working with the crypto module and you start four password-encryption commands in a row. How long will each command take, and will the times differ significantly?",
        [
            "Ответ зависит от того, используются синхронные или асинхронные crypto-операции. Если вызвать синхронные функции вроде pbkdf2Sync или bcrypt sync-variant, они выполняются последовательно и блокируют event loop. В этом случае каждая команда будет ждать завершения предыдущей, а приложение не сможет нормально обрабатывать другие запросы.",
            "Если используются асинхронные операции Node.js crypto, тяжелая работа обычно уходит в libuv thread pool. По умолчанию размер пула — 4 потока. Если запустить ровно 4 тяжелые операции подряд, они могут стартовать параллельно и завершиться примерно в одно и то же время, если параметры одинаковые и система не перегружена. Разница будет, но обычно не драматическая: планирование ОС, нагрузка CPU и другие задачи могут добавить разброс.",
            "Если запустить пятую такую операцию, она будет ждать свободный поток в пуле. Поэтому первые четыре могут завершиться примерно через T, а пятая — примерно через 2T, если все задачи одинаковой тяжести. Размер пула можно изменить через UV_THREADPOOL_SIZE, но это нужно делать осторожно: больше потоков не всегда быстрее, потому что CPU может стать bottleneck.",
            ("bullets", [
                "Sync crypto блокирует event loop и выполняется последовательно.",
                "Async crypto использует libuv thread pool для тяжелых операций.",
                "По умолчанию в пуле 4 потока, поэтому 4 задачи могут выполняться параллельно.",
                "Пятая задача обычно ждет, пока освободится один из потоков.",
                "Время выполнения зависит от CPU, параметров алгоритма и общей нагрузки процесса."
            ]),
        ],
        [
            "The answer depends on whether synchronous or asynchronous crypto operations are used. If synchronous functions such as pbkdf2Sync or a synchronous bcrypt variant are called, they run sequentially and block the event loop. In that case, each command waits for the previous one to finish, and the application cannot properly handle other requests during that time.",
            "If asynchronous Node.js crypto operations are used, the heavy work usually runs in the libuv thread pool. The default pool size is 4. If exactly four heavy operations are started in a row, they can run in parallel and finish at roughly the same time if their parameters are identical and the system is not overloaded. The timings may differ, but usually not dramatically; OS scheduling, CPU load, and other queued work can add variance.",
            "If a fifth operation is started, it usually waits for a free thread in the pool. Therefore, the first four tasks may finish after roughly T, while the fifth may finish around 2T if all tasks have equal cost. The pool size can be changed through UV_THREADPOOL_SIZE, but this should be done carefully: more threads are not always faster if the CPU becomes the bottleneck.",
            ("bullets", [
                "Synchronous crypto blocks the event loop and runs sequentially.",
                "Asynchronous crypto uses the libuv thread pool for heavy work.",
                "The default pool size is 4, so four tasks can run in parallel.",
                "A fifth task usually waits until one thread becomes free.",
                "Execution time depends on CPU, algorithm parameters, and process load."
            ]),
        ],
    )

    add_section(doc, "AI Content Platform")

    add_qa(
        doc,
        9,
        "Пользователь загружает двухчасовое видео. Как ты организуешь обработку файла в Node.js, чтобы не положить память процесса?\nПроверяет: streams, backpressure, event loop, S3/MinIO.",
        "A user uploads a two-hour video. How would you organize file processing in Node.js so the process memory does not collapse?\nChecks: streams, backpressure, event loop, S3/MinIO.",
        [
            "Главный принцип — не читать видео целиком в память. Двухчасовой файл может быть сотни мегабайт или несколько гигабайт, поэтому Buffer всего файла легко приведет к out-of-memory или сильному pressure на garbage collector. Нужно использовать streaming upload или, еще лучше, direct upload в S3/MinIO через signed URL.",
            "Если backend принимает файл, он должен работать с потоками: request stream передается дальше в storage stream или multipart upload. Backpressure важен, потому что consumer потока может быть медленнее producer. Правильная stream pipeline автоматически замедляет чтение, если запись в S3, диск или другой stream не успевает.",
            "После загрузки backend не должен синхронно обрабатывать видео в request/response цикле. Нужно создать job в PostgreSQL, положить событие в Kafka или очередь и вернуть клиенту jobId. Дальше worker обрабатывает видео асинхронно, обновляет status/progress и сохраняет результаты в S3/DB.",
            ("bullets", [
                "Ограничить размер файла и тип MIME.",
                "Использовать multipart upload и signed URLs, если возможно.",
                "Не хранить весь файл в Buffer.",
                "Создавать async job после загрузки.",
                "Обновлять progress через DB/Redis/WebSocket, но source of truth держать в БД."
            ]),
        ],
        [
            "The main principle is to never read the whole video into memory. A two-hour file can be hundreds of megabytes or several gigabytes, so buffering the entire file can easily cause out-of-memory errors or heavy garbage collector pressure. The system should use streaming upload or, even better, direct upload to S3/MinIO through a signed URL.",
            "If the backend receives the file, it should work with streams: the request stream is piped to a storage stream or multipart upload. Backpressure matters because the stream consumer can be slower than the producer. A proper stream pipeline automatically slows down reading when writing to S3, disk, or another stream cannot keep up.",
            "After upload, the backend should not process the video synchronously inside the request/response cycle. It should create a job in PostgreSQL, publish an event to Kafka or another queue, and return a jobId to the client. A worker then processes the video asynchronously, updates status/progress, and stores results in S3/DB.",
            ("bullets", [
                "Limit file size and MIME type.",
                "Use multipart upload and signed URLs when possible.",
                "Never keep the whole file in a Buffer.",
                "Create an async job after upload.",
                "Update progress through DB/Redis/WebSocket, while keeping the database as the source of truth."
            ]),
        ],
    )

    add_qa(
        doc,
        10,
        "В системе есть этапы: upload, transcription, summarization, segmentation, publishing. Как бы ты описал их типами TypeScript, чтобы нельзя было случайно обработать job в неверном состоянии?\nПроверяет: discriminated unions,\u00a0never, strict typing.",
        "The system has stages: upload, transcription, summarization, segmentation, publishing. How would you describe them with TypeScript types so a job cannot accidentally be processed in the wrong state?\nChecks: discriminated unions, never, strict typing.",
        [
            "Лучший подход — использовать discriminated union, где поле status или stage является дискриминатором. Для каждого состояния описывается свой набор обязательных и допустимых полей. Например, uploaded может иметь mediaUrl, transcribing — transcriptionJobId, summarized — summary, failed — error, completed — outputIds.",
            "Тогда TypeScript сможет сужать тип внутри switch по status. Если разработчик пытается прочитать summary у job в состоянии uploaded, компилятор покажет ошибку. Это снижает вероятность runtime-багов, особенно в long-running workflows, где состояние меняется постепенно.",
            "Дополнительно полезно использовать exhaustive checking через never. В default ветке switch можно присвоить значение переменной типа never. Если позже добавится новый status, но обработчик не обновится, TypeScript начнет ругаться. Это особенно важно в системах с несколькими worker-ами и event handlers.",
            ("bullets", [
                "status/stage должен быть явным discriminator.",
                "Каждый status имеет собственный набор полей.",
                "Переходы между состояниями лучше оформлять отдельными функциями.",
                "Exhaustive switch через never ловит забытые новые состояния.",
                "Runtime validation все равно нужна для данных из Kafka/API."
            ]),
        ],
        [
            "The best approach is to use a discriminated union where status or stage acts as the discriminator. Each state has its own required and allowed fields. For example, uploaded can have mediaUrl, transcribing can have transcriptionJobId, summarized can have summary, failed can have error, and completed can have outputIds.",
            "TypeScript can then narrow the type inside a switch by status. If a developer tries to read summary from a job in the uploaded state, the compiler will report an error. This reduces runtime bugs, especially in long-running workflows where state changes step by step.",
            "It is also useful to add exhaustive checking with never. In the default branch of a switch, assign the value to a variable of type never. If a new status is added later and the handler is not updated, TypeScript will fail compilation. This is especially important in systems with multiple workers and event handlers.",
            ("bullets", [
                "status/stage should be an explicit discriminator.",
                "Each status should define its own fields.",
                "State transitions are best represented by dedicated functions.",
                "An exhaustive switch with never catches newly added states.",
                "Runtime validation is still required for data coming from Kafka/API."
            ]),
        ],
    )

    add_qa(
        doc,
        11,
        "Kafka consumer получил задачу генерации, вызвал OpenAI API, но упал до сохранения результата в PostgreSQL. Как система должна восстановиться?\nПроверяет: Kafka offsets, retries, idempotency, transactions.",
        "A Kafka consumer received a generation task, called the OpenAI API, but crashed before saving the result to PostgreSQL. How should the system recover?\nChecks: Kafka offsets, retries, idempotency, transactions.",
        [
            "Consumer не должен коммитить offset до того, как результат надежно сохранен или задача переведена в понятное terminal state. Если процесс упал до commit, Kafka доставит сообщение повторно. Поэтому обработчик обязан быть идемпотентным: повторная обработка одного и того же jobId не должна создавать дубликаты результатов или портить статус.",
            "Практический вариант — хранить state machine в PostgreSQL: jobId, currentStep, attempt, externalRequestId, resultHash, status. Перед вызовом OpenAI можно пометить step как in_progress, а после получения результата сохранить его с уникальным ключом jobId + step. Если retry пришел повторно, worker проверяет, был ли step уже завершен.",
            "Если результат сохраняется в S3, object key лучше делать детерминированным, например по jobId и output type. Тогда повторная запись не создаст хаотичные дубликаты. Для сложных систем полезны outbox/inbox pattern, DLQ для terminal failures и backoff retries для transient failures.",
            ("bullets", [
                "Offset коммитится после надежного сохранения результата.",
                "Обработка должна быть идемпотентной.",
                "Уникальные constraints помогают избежать дубликатов.",
                "Transient errors ретраятся, terminal errors уходят в failed/DLQ.",
                "Exactly-once с внешним OpenAI практически недостижим, поэтому нужна at-least-once + deduplication."
            ]),
        ],
        [
            "The consumer should not commit the offset until the result is reliably saved or the task is moved into a clear terminal state. If the process crashes before committing, Kafka will deliver the message again. Therefore, the handler must be idempotent: processing the same jobId again must not create duplicate results or corrupt the status.",
            "A practical approach is to store a state machine in PostgreSQL: jobId, currentStep, attempt, externalRequestId, resultHash, and status. Before calling OpenAI, the worker can mark the step as in_progress. After receiving the result, it saves it with a unique key such as jobId + step. If a retry arrives, the worker checks whether the step has already been completed.",
            "If the result is stored in S3, the object key should be deterministic, for example based on jobId and output type. Then a retry does not create random duplicates. For complex systems, outbox/inbox patterns, DLQ for terminal failures, and backoff retries for transient failures are useful.",
            ("bullets", [
                "The offset is committed only after reliable persistence.",
                "The handler must be idempotent.",
                "Unique constraints help prevent duplicates.",
                "Transient errors are retried; terminal errors go to failed/DLQ.",
                "Exactly-once with an external OpenAI call is practically impossible, so use at-least-once plus deduplication."
            ]),
        ],
    )

    add_qa(
        doc,
        12,
        "Почему для AI-задач Kafka уместнее, чем прямой HTTP-вызов между сервисами? А где HTTP всё равно лучше?\nПроверяет: async architecture, microservices vs simple calls.",
        "Why is Kafka more appropriate than a direct HTTP call between services for AI tasks? Where is HTTP still better?\nChecks: async architecture, microservices vs simple calls.",
        [
            "AI-задачи часто долгие, дорогие и нестабильные по времени: транскрибация, суммаризация и генерация могут занимать минуты. HTTP-вызов в таком случае держал бы соединение открытым, был бы чувствителен к timeout и плохо переносил бы временные сбои downstream-сервиса. Kafka лучше подходит как буфер и механизм асинхронной доставки.",
            "Kafka позволяет развязать producer и consumer: API service быстро создает job, публикует событие и возвращает клиенту jobId. Worker обрабатывает задачу в своем темпе. Очередь помогает с retries, backpressure, масштабированием consumers и устойчивостью к временным падениям processing service.",
            "Но Kafka не нужна везде. HTTP лучше для быстрых синхронных операций, где пользователь ждет немедленный ответ: получить статус job, загрузить список кампаний, обновить профиль, проверить права доступа. Если операция проста, быстра и не требует асинхронной доставки, Kafka только усложнит систему.",
            ("bullets", [
                "Kafka хороша для долгих задач, очередей, retries и backpressure.",
                "HTTP хорош для быстрых request/response операций.",
                "Kafka не заменяет API, а дополняет его.",
                "Выбор зависит от latency, reliability и необходимости асинхронной обработки."
            ]),
        ],
        [
            "AI tasks are often long-running, expensive, and unpredictable in duration: transcription, summarization, and generation may take minutes. An HTTP call would keep the connection open, be sensitive to timeouts, and handle downstream service failures poorly. Kafka is better suited as a buffer and asynchronous delivery mechanism.",
            "Kafka decouples the producer and consumer: the API service quickly creates a job, publishes an event, and returns a jobId to the client. The worker processes the task at its own pace. The queue helps with retries, backpressure, consumer scaling, and resilience when the processing service temporarily fails.",
            "However, Kafka is not needed everywhere. HTTP is better for fast synchronous operations where the user expects an immediate response: reading job status, listing campaigns, updating a profile, or checking permissions. If the operation is simple, fast, and does not require asynchronous delivery, Kafka only adds complexity.",
            ("bullets", [
                "Kafka is strong for long-running jobs, queues, retries, and backpressure.",
                "HTTP is strong for fast request/response operations.",
                "Kafka does not replace an API; it complements it.",
                "The choice depends on latency, reliability, and the need for asynchronous processing."
            ]),
        ],
    )

    add_qa(
        doc,
        13,
        "Пользователь видит progress через WebSocket. Что произойдет, если он обновит страницу, а worker продолжает обработку?\nПроверяет: WebSockets, source of truth, reconnect.",
        "The user sees progress through WebSocket. What happens if the user refreshes the page while the worker continues processing?\nChecks: WebSockets, source of truth, reconnect.",
        [
            "WebSocket-соединение разорвется, но job не должен зависеть от этого соединения. Worker продолжает обработку, потому что source of truth находится не в WebSocket, а в устойчивом хранилище: PostgreSQL для статуса job и, возможно, Redis для transient progress или pub/sub.",
            "После обновления страницы frontend должен заново получить актуальный snapshot через REST или server-rendered data: текущий status, процент выполнения, последний завершенный step, ошибки и результаты. После этого клиент снова подписывается на WebSocket events по jobId или campaignId.",
            "Если полагаться только на события WebSocket, пользователь может потерять часть обновлений во время reconnect. Поэтому события должны быть дополнением к состоянию, а не единственным источником истины. В более надежной системе события имеют sequence number или timestamp, чтобы клиент мог понять, какие updates он пропустил.",
            ("bullets", [
                "WebSocket — транспорт уведомлений, не source of truth.",
                "После refresh клиент делает REST-запрос за snapshot.",
                "Worker продолжает job независимо от подключения пользователя.",
                "Для надежности можно использовать sequence numbers и replay последних events."
            ]),
        ],
        [
            "The WebSocket connection is closed, but the job should not depend on that connection. The worker continues processing because the source of truth is not the WebSocket, but durable storage: PostgreSQL for job status and possibly Redis for transient progress or pub/sub.",
            "After the page refresh, the frontend should fetch the current snapshot again through REST or server-rendered data: current status, progress percentage, last completed step, errors, and results. Then the client subscribes to WebSocket events again by jobId or campaignId.",
            "If the system relies only on WebSocket events, the user may miss updates during reconnect. Therefore, events should complement state, not be the only source of truth. In a more reliable system, events have a sequence number or timestamp so the client can detect missed updates.",
            ("bullets", [
                "WebSocket is a notification transport, not the source of truth.",
                "After refresh, the client fetches a snapshot through REST.",
                "The worker continues the job independently of the user's connection.",
                "For reliability, use sequence numbers and optional replay of recent events."
            ]),
        ],
    )

    add_qa(
        doc,
        14,
        "Если WebSocket-сервис запущен в нескольких ECS tasks, как доставить progress event именно нужному пользователю?\nПроверяет: scaling, pub/sub, sticky sessions.",
        "If the WebSocket service is running in several ECS tasks, how do you deliver a progress event to the correct user?\nChecks: scaling, pub/sub, sticky sessions.",
        [
            "При нескольких ECS tasks проблема в том, что пользователь может быть подключен к одному instance, а worker или API service публикует событие в другой instance. Локальная in-memory карта socketId -> userId работает только внутри одного процесса и ломается при горизонтальном масштабировании.",
            "Нужен общий слой доставки событий. Частый вариант — Redis Pub/Sub или Redis adapter для socket.io: каждый instance подписан на общий канал, а событие публикуется с jobId/userId. Instance, у которого есть нужное соединение, отправляет сообщение клиенту. Альтернатива — Kafka/NATS topic или отдельный notification service, который знает активные подключения.",
            "Sticky sessions могут помочь удерживать пользователя на одном instance, но они не решают всю проблему: событие все равно может прийти не туда. Поэтому sticky sessions — это оптимизация соединений, а не полноценный механизм routing событий. Также нужно проверять авторизацию подписки: пользователь не должен подписаться на чужой jobId.",
            ("bullets", [
                "In-memory socket registry не масштабируется между ECS tasks.",
                "Нужен общий pub/sub или централизованный notification service.",
                "Sticky sessions помогают, но не заменяют межпроцессную доставку.",
                "Событие должно содержать jobId/userId/tenantId для routing и authz."
            ]),
        ],
        [
            "With multiple ECS tasks, the issue is that the user may be connected to one instance while the worker or API service publishes the event to another. A local in-memory map such as socketId -> userId only works inside one process and breaks under horizontal scaling.",
            "A shared event delivery layer is needed. A common option is Redis Pub/Sub or a Redis adapter for socket.io: every instance subscribes to a shared channel, and events are published with jobId/userId. The instance that owns the relevant connection sends the message to the client. Alternatives include a Kafka/NATS topic or a dedicated notification service that tracks active connections.",
            "Sticky sessions can help keep a user connected to the same instance, but they do not solve the full problem: an event can still arrive at a different instance. Therefore, sticky sessions are a connection optimization, not a complete event routing mechanism. Authorization must also be checked: a user must not be able to subscribe to someone else's jobId.",
            ("bullets", [
                "An in-memory socket registry does not scale across ECS tasks.",
                "Use a shared pub/sub layer or centralized notification service.",
                "Sticky sessions help, but do not replace cross-process delivery.",
                "Events should include jobId/userId/tenantId for routing and authorization."
            ]),
        ],
    )

    add_qa(
        doc,
        15,
        "Как бы ты ограничил стоимость OpenAI API на пользователя, кампанию и организацию?\nПроверяет: rate limiting, quotas, backend design, observability.",
        "How would you limit OpenAI API cost per user, campaign, and organization?\nChecks: rate limiting, quotas, backend design, observability.",
        [
            "Стоимость нужно контролировать не только rate limit-ами, но и бюджетами. Для каждого запроса нужно считать model, input tokens, output tokens, estimated cost, userId, organizationId, campaignId и jobId. Эти данные стоит сохранять как usage events, чтобы потом строить лимиты, аналитику и алерты.",
            "На уровне backend можно ввести quotas: дневной/месячный лимит на организацию, лимит на кампанию, лимит на размер transcript, лимит concurrent jobs. Перед созданием generation job система проверяет доступный бюджет. Если лимит исчерпан, job не создается или ставится в paused state до подтверждения.",
            "На уровне архитектуры полезны rate limiter, concurrency limiter и model routing. Например, предварительную классификацию можно делать дешевой моделью, а дорогую модель использовать только для финального quality pass. Повторяемые операции можно кэшировать по input hash + prompt version + model settings.",
            ("bullets", [
                "Считать cost и tokens на каждом LLM-вызове.",
                "Ввести quotas по user/org/campaign.",
                "Ограничивать размер входных данных и параллельность jobs.",
                "Использовать дешевые модели для простых этапов.",
                "Добавить budget alerts и dashboards."
            ]),
        ],
        [
            "Cost should be controlled not only with rate limits, but also with budgets. For each request, the system should record model, input tokens, output tokens, estimated cost, userId, organizationId, campaignId, and jobId. These should be stored as usage events so limits, analytics, and alerts can be built on top.",
            "At the backend level, quotas can be introduced: daily/monthly organization limits, campaign limits, transcript size limits, and concurrent job limits. Before creating a generation job, the system checks the available budget. If the limit is exhausted, the job is not created or is placed into a paused state until confirmed.",
            "Architecturally, rate limiters, concurrency limiters, and model routing are useful. For example, preliminary classification can use a cheaper model, while an expensive model is reserved for the final quality pass. Repeated operations can be cached by input hash + prompt version + model settings.",
            ("bullets", [
                "Track cost and tokens for every LLM call.",
                "Introduce quotas by user/org/campaign.",
                "Limit input size and concurrent jobs.",
                "Use cheaper models for simpler stages.",
                "Add budget alerts and dashboards."
            ]),
        ],
    )

    add_qa(
        doc,
        16,
        "OpenAI вернул невалидный JSON. Где должна быть валидация: в LangChain, service layer, DTO, frontend?\nПроверяет: validation, error handling, schema checks.",
        "OpenAI returned invalid JSON. Where should validation happen: in LangChain, service layer, DTO, or frontend?\nChecks: validation, error handling, schema checks.",
        [
            "Главная валидация должна быть в backend service layer сразу после получения LLM output. Frontend не должен быть первым местом, где обнаруживается невалидный JSON, потому что к этому моменту результат уже мог быть сохранен, показан пользователю или использован в следующем шаге pipeline.",
            "LangChain или structured output parser можно использовать как первый слой: попросить модель вернуть JSON по схеме, применить parser, retry или repair strategy. Но это не отменяет собственной backend-валидации. DTO часто валидирует входящие HTTP-запросы, а здесь нужно валидировать внешний ответ LLM, поэтому лучше использовать schema validation: Zod, JSON Schema, class-validator или аналог.",
            "Если JSON невалидный, система должна решить: это transient ошибка, которую можно повторить, или terminal failure. Возможные стратегии: retry с уточненным prompt, JSON repair, fallback model, сохранение failed status с причиной. В любом случае нельзя молча принимать частично неправильную структуру.",
            ("bullets", [
                "LangChain/parser помогает, но не является единственной защитой.",
                "Service layer должен валидировать LLM output перед сохранением.",
                "Frontend получает уже нормализованный и валидный результат.",
                "Ошибки нужно логировать с jobId, promptVersion и model.",
                "Для стабильности полезны schema validation и retry policy."
            ]),
        ],
        [
            "The main validation should happen in the backend service layer immediately after receiving the LLM output. The frontend should not be the first place where invalid JSON is discovered, because by that time the result may already have been saved, shown to the user, or used in the next pipeline step.",
            "LangChain or a structured output parser can be used as the first layer: ask the model to return JSON according to a schema, apply a parser, retry, or repair strategy. But this does not replace backend validation. DTOs often validate incoming HTTP requests, while here the system must validate an external LLM response, so schema validation with Zod, JSON Schema, class-validator, or a similar tool is better.",
            "If the JSON is invalid, the system must decide whether this is a transient error that can be retried or a terminal failure. Possible strategies include retry with a stricter prompt, JSON repair, fallback model, or saving a failed status with a clear reason. The system must not silently accept a partially incorrect structure.",
            ("bullets", [
                "LangChain/parser helps, but is not the only protection.",
                "The service layer must validate LLM output before persistence.",
                "The frontend should receive normalized and valid data.",
                "Errors should be logged with jobId, promptVersion, and model.",
                "Schema validation and retry policy improve stability."
            ]),
        ],
    )

    add_section(doc, "Node.js / Backend")

    add_qa(
        doc,
        17,
        "Node.js / Backend\n9. В AI processing service одновременно идут 100 тяжелых задач. CPU низкий, latency высокий. Какие гипотезы проверишь?\nПроверяет: event loop, libuv, connection pools, external API latency.",
        "Node.js / Backend\n9. In an AI processing service, 100 heavy tasks are running at the same time. CPU is low, latency is high. What hypotheses would you check?\nChecks: event loop, libuv, connection pools, external API latency.",
        [
            "Низкий CPU при высокой latency часто означает, что процесс не считает, а ждет. Возможные bottlenecks: OpenAI API latency, S3 upload/download, PostgreSQL connection pool, Redis latency, Kafka consumer lag, DNS/TLS handshake, network saturation или слишком много одновременных Promise, ожидающих внешние ресурсы.",
            "Я бы начал с измерений по этапам job: время ожидания в очереди, время запроса к OpenAI, время записи в БД, время загрузки в S3, время публикации progress events. Нужны structured logs с correlationId/jobId и метрики p50/p95/p99. Также стоит смотреть event loop delay: если он высокий, значит main thread блокируется CPU-bound задачей или тяжелой синхронной операцией.",
            "Дальше проверяются лимиты: размер DB pool, лимиты OpenAI, количество concurrent workers, размер libuv thread pool, Kafka lag, retries. Если одновременно запущено 100 задач без ограничителя, система может создать очередь ожидания на внешнем API или пуле соединений. Поэтому нужен concurrency limiter и backpressure на уровне worker-а.",
            ("bullets", [
                "Низкий CPU + высокая latency часто указывает на IO wait.",
                "Нужен breakdown latency по каждому этапу pipeline.",
                "Проверить DB/Redis/S3/OpenAI/Kafka, а не только Node.js.",
                "Добавить concurrency limits, timeouts и retries with backoff.",
                "Следить за event loop delay и размером очередей."
            ]),
        ],
        [
            "Low CPU with high latency often means the process is not computing, but waiting. Possible bottlenecks include OpenAI API latency, S3 upload/download, PostgreSQL connection pool, Redis latency, Kafka consumer lag, DNS/TLS handshake, network saturation, or too many concurrent Promises waiting on external resources.",
            "I would start by measuring each job stage: queue wait time, OpenAI request time, database write time, S3 transfer time, and progress event publishing time. Structured logs with correlationId/jobId and p50/p95/p99 metrics are needed. Event loop delay should also be checked: if it is high, the main thread may be blocked by CPU-bound work or heavy synchronous operations.",
            "Then I would check limits: DB pool size, OpenAI limits, worker concurrency, libuv thread pool size, Kafka lag, and retries. If 100 tasks are started without a limiter, the system may build an internal queue at the external API or connection pool level. A worker-level concurrency limiter and backpressure are required.",
            ("bullets", [
                "Low CPU + high latency often points to IO wait.",
                "Latency breakdown is needed for each pipeline stage.",
                "Check DB/Redis/S3/OpenAI/Kafka, not only Node.js.",
                "Add concurrency limits, timeouts, and retries with backoff.",
                "Monitor event loop delay and queue sizes."
            ]),
        ],
    )

    add_qa(
        doc,
        18,
        "10. Когда в Node.js стоит использовать worker threads, а когда они не помогут? Приведи пример из video/content processing.\nПроверяет: CPU-bound vs IO-bound.",
        "10. When should worker threads be used in Node.js, and when will they not help? Give an example from video/content processing.\nChecks: CPU-bound vs IO-bound.",
        [
            "Worker threads нужны, когда задача CPU-bound: тяжелые вычисления занимают main thread и блокируют event loop. Примеры: парсинг большого файла, компрессия, криптографические операции, генерация thumbnails, обработка изображений, тяжелый анализ текста или локальная обработка видео. В таких случаях worker thread позволяет вынести вычисление из основного потока и сохранить отзывчивость сервера.",
            "Worker threads не помогут, если bottleneck — ожидание внешнего ресурса: OpenAI API, PostgreSQL, Redis, Kafka, S3 или сеть. Если задача большую часть времени ждет ответ API, перенос ее в worker только добавит overhead на передачу данных между потоками. Для IO-bound задач лучше использовать async/await, connection pools, retries, timeout и concurrency limits.",
            "В video/content processing важно отделить orchestration от тяжелой обработки. Например, если сервис просто отправляет файл в cloud transcription API и ждет результат, worker thread не нужен. Если сервис локально режет видео, извлекает кадры или сжимает данные, worker thread или отдельный processing service будет уместен.",
            ("bullets", [
                "CPU-bound задачи можно выносить в worker threads.",
                "IO-bound задачи лучше решать асинхронностью и лимитами параллельности.",
                "Передача больших объектов между потоками может быть дорогой.",
                "Для очень тяжелой обработки иногда лучше отдельный сервис, а не worker в том же процессе."
            ]),
        ],
        [
            "Worker threads are useful when a task is CPU-bound: heavy computation occupies the main thread and blocks the event loop. Examples include parsing a large file, compression, cryptographic operations, thumbnail generation, image processing, heavy text analysis, or local video processing. In such cases, a worker thread moves computation away from the main thread and keeps the server responsive.",
            "Worker threads do not help when the bottleneck is waiting for an external resource: OpenAI API, PostgreSQL, Redis, Kafka, S3, or the network. If a task spends most of its time waiting for an API response, moving it to a worker only adds overhead for passing data between threads. For IO-bound tasks, async/await, connection pools, retries, timeouts, and concurrency limits are better tools.",
            "In video/content processing, orchestration should be separated from heavy processing. If the service simply sends a file to a cloud transcription API and waits for the result, a worker thread is not needed. If the service locally cuts video, extracts frames, or compresses data, a worker thread or separate processing service is appropriate.",
            ("bullets", [
                "CPU-bound work can be moved to worker threads.",
                "IO-bound work is better handled with async flow and concurrency limits.",
                "Passing large objects between threads can be expensive.",
                "For very heavy processing, a separate service may be better than a worker inside the same process."
            ]),
        ],
    )

    add_section(doc, "Databases")

    add_qa(
        doc,
        19,
        "Databases\n11. Для AI campaign storage: что хранишь в PostgreSQL, что в S3, что в Redis? Почему?\nПроверяет: SQL, object storage, cache, source of truth.",
        "Databases\n11. For AI campaign storage, what would you store in PostgreSQL, what in S3, and what in Redis? Why?\nChecks: SQL, object storage, cache, source of truth.",
        [
            "PostgreSQL я бы использовал как source of truth для бизнес-сущностей: users, organizations, campaigns, media metadata, generation_jobs, statuses, prompt versions, usage/cost records, audit log. Эти данные требуют консистентности, индексов, транзакций и понятной модели отношений.",
            "S3 или MinIO подходит для больших бинарных объектов: raw video, audio, thumbnails, generated video snippets, exported documents, возможно transcript files, если они большие. В PostgreSQL лучше хранить не сам файл, а metadata и object key: bucket, path, size, checksum, contentType, ownerId, createdAt.",
            "Redis я бы использовал для временных и ускоряющих данных: cache, rate limits, distributed locks, pub/sub для WebSocket notifications, transient progress или short-lived session-like state. Но Redis не должен быть единственным хранилищем статуса job, потому что ключ может истечь или быть потерян при сбое.",
            ("bullets", [
                "PostgreSQL: бизнес-данные, статусы, audit, usage.",
                "S3/MinIO: большие файлы и бинарные результаты.",
                "Redis: cache, rate limits, pub/sub, ephemeral progress.",
                "Source of truth для job status лучше держать в PostgreSQL."
            ]),
        ],
        [
            "I would use PostgreSQL as the source of truth for business entities: users, organizations, campaigns, media metadata, generation_jobs, statuses, prompt versions, usage/cost records, and audit logs. These data require consistency, indexes, transactions, and a clear relational model.",
            "S3 or MinIO is appropriate for large binary objects: raw video, audio, thumbnails, generated video snippets, exported documents, and possibly transcript files if they are large. PostgreSQL should store metadata and the object key rather than the file itself: bucket, path, size, checksum, contentType, ownerId, and createdAt.",
            "Redis should be used for temporary and performance-oriented data: cache, rate limits, distributed locks, pub/sub for WebSocket notifications, transient progress, or short-lived session-like state. But Redis should not be the only storage for job status because keys can expire or be lost during failures.",
            ("bullets", [
                "PostgreSQL: business data, statuses, audit, usage.",
                "S3/MinIO: large files and binary results.",
                "Redis: cache, rate limits, pub/sub, ephemeral progress.",
                "The source of truth for job status should live in PostgreSQL."
            ]),
        ],
    )

    add_section(doc, "Frontend State and Rendering")

    add_qa(
        doc,
        20,
        "Frontend / React / Next.js\n12. На странице AI generation есть upload progress, job status, WebSocket events, generated results и ошибки. Где хранить каждое состояние?\nПроверяет: React state, Redux, local/server state.",
        "Frontend / React / Next.js\n12. On an AI generation page there are upload progress, job status, WebSocket events, generated results, and errors. Where would you store each piece of state?\nChecks: React state, Redux, local/server state.",
        [
            "Эти состояния лучше разделить по природе данных. Upload progress — локальное transient state конкретного компонента загрузки, потому что оно живет только во время текущей загрузки. Его можно хранить в useState/useReducer или в state upload hook-а. Глобальный store для этого обычно не нужен.",
            "Job status и generated results — server state. Их источник истины находится на backend, поэтому удобно использовать SWR, React Query, server components или обычный fetch + cache strategy. WebSocket events не должны полностью заменять server state: они применяются как updates к snapshot, полученному с backend.",
            "Ошибки нужно разделять: upload error, API error, WebSocket reconnect error, terminal job error. Recoverable ошибки можно показывать как notification/retry state, terminal ошибки — как часть job status. Redux имеет смысл только если эти данные нужны нескольким далеким частям приложения: например, global notifications, user session, organization context или cross-page job indicator.",
            ("bullets", [
                "Upload progress: local component state.",
                "Job status: server state с периодическим refetch/snapshot.",
                "WebSocket events: incremental updates, не source of truth.",
                "Generated results: server state/cache.",
                "Errors: разделять по источнику и recoverability.",
                "Redux: только для действительно глобального состояния."
            ]),
        ],
        [
            "These states should be separated by their nature. Upload progress is local transient state for the upload component because it exists only during the current upload. It can be stored in useState/useReducer or in a dedicated upload hook. A global store is usually unnecessary for this.",
            "Job status and generated results are server state. Their source of truth is on the backend, so SWR, React Query, server components, or regular fetch with a cache strategy are appropriate. WebSocket events should not fully replace server state; they should be applied as updates to a backend snapshot.",
            "Errors should be separated: upload error, API error, WebSocket reconnect error, and terminal job error. Recoverable errors can be shown as notifications/retry state, while terminal errors should be part of the job status. Redux makes sense only if the data is needed by several distant parts of the application: global notifications, user session, organization context, or a cross-page job indicator.",
            ("bullets", [
                "Upload progress: local component state.",
                "Job status: server state with periodic refetch/snapshot.",
                "WebSocket events: incremental updates, not source of truth.",
                "Generated results: server state/cache.",
                "Errors: separate by source and recoverability.",
                "Redux: only for truly global state."
            ]),
        ],
    )

    add_qa(
        doc,
        21,
        "13 Какие данные AI Content Platform ты бы рендерил через SSR/Next.js, а какие только на клиенте?\nПроверяет: SSR, CSR, Next.js trade-offs.",
        "13 Which AI Content Platform data would you render through SSR/Next.js, and which only on the client?\nChecks: SSR, CSR, Next.js trade-offs.",
        [
            "Через SSR/Server Components я бы рендерил данные, которые нужны для первого meaningful paint и не требуют постоянной интерактивности: layout, user/org context, список кампаний, metadata кампании, начальный snapshot job status, результаты, которые уже готовы. Это помогает быстрее показать страницу и уменьшить клиентский JavaScript.",
            "Только на клиенте лучше держать интерактивные и быстро меняющиеся части: upload progress, drag-and-drop file input, WebSocket connection, optimistic UI, live progress updates, локальные фильтры и действия пользователя. Эти части требуют browser APIs и должны быть client components.",
            "Для Next.js важно не превращать всю страницу в 'use client'. Лучше сделать страницу server component, загрузить начальные данные на сервере, а внутрь вставить небольшие client components для upload/progress/actions. Так уменьшается hydration cost и размер client bundle.",
            ("bullets", [
                "SSR/Server Components: initial data, metadata, уже готовые результаты.",
                "Client Components: upload, WebSocket, live progress, interactive controls.",
                "CSR полезен для динамики, но не должен забирать всю страницу без причины.",
                "Граница server/client должна проходить по необходимости интерактивности."
            ]),
        ],
        [
            "I would render data needed for the first meaningful paint through SSR/Server Components when it does not require constant interactivity: layout, user/org context, campaign list, campaign metadata, initial job status snapshot, and already completed results. This helps show the page faster and reduce client-side JavaScript.",
            "Client-side rendering is better for interactive and rapidly changing parts: upload progress, drag-and-drop file input, WebSocket connection, optimistic UI, live progress updates, local filters, and user actions. These parts require browser APIs and should be client components.",
            "In Next.js, it is important not to turn the whole page into 'use client'. A better approach is to keep the page as a server component, load initial data on the server, and place small client components inside for upload/progress/actions. This reduces hydration cost and client bundle size.",
            ("bullets", [
                "SSR/Server Components: initial data, metadata, completed results.",
                "Client Components: upload, WebSocket, live progress, interactive controls.",
                "CSR is useful for dynamic behavior, but should not take over the whole page without reason.",
                "The server/client boundary should follow the need for interactivity."
            ]),
        ],
    )

    add_section(doc, "Security / Web")

    add_qa(
        doc,
        22,
        "Security / Web\n14. В резюме есть OAuth 2.0, JWT и Passport.js. Где в его проектах логично использовать OAuth, а где обычную JWT-сессию?\nПроверяет: auth concepts, practical boundaries.",
        "Security / Web\n14. The CV mentions OAuth 2.0, JWT, and Passport.js. Where in his projects would OAuth be logical, and where would a regular JWT session be enough?\nChecks: auth concepts, practical boundaries.",
        [
            "OAuth 2.0 логичен, когда пользователь входит через внешнего провайдера или приложение получает делегированный доступ к внешнему ресурсу: Google, GitHub, Facebook, YouTube, social media publishing APIs. OAuth решает задачу authorization delegation: пользователь разрешает приложению действовать в определенных рамках без передачи пароля.",
            "JWT-сессия уместна для доступа к собственному backend API после того, как пользователь уже прошел authentication. Например, после OAuth login backend может создать внутреннего пользователя и выдать access token/refresh token для запросов к своим сервисам. То есть OAuth может быть внешним механизмом входа, а JWT — внутренним механизмом авторизованных запросов.",
            "Passport.js в Node/Nest/Express часто используется как набор стратегий: local strategy для email/password, jwt strategy для проверки access token, oauth strategies для внешних провайдеров. Важно не смешивать понятия: OAuth — протокол делегированной авторизации, JWT — формат токена, Passport — middleware/framework для реализации стратегий.",
            ("bullets", [
                "OAuth: внешний провайдер или delegated access.",
                "JWT: внутренний access token к собственному API.",
                "Refresh token лучше хранить безопаснее, чем обычный JS-accessible localStorage.",
                "Authz checks все равно должны выполняться на backend/service layer."
            ]),
        ],
        [
            "OAuth 2.0 is logical when the user signs in through an external provider or when the application receives delegated access to an external resource: Google, GitHub, Facebook, YouTube, or social media publishing APIs. OAuth solves authorization delegation: the user allows the application to act within a defined scope without sharing a password.",
            "A JWT session is appropriate for accessing the application's own backend API after the user has already been authenticated. For example, after OAuth login, the backend can create an internal user and issue an access token/refresh token for requests to internal services. In other words, OAuth can be the external login mechanism, while JWT is the internal mechanism for authorized API requests.",
            "Passport.js in Node/Nest/Express is often used as a set of strategies: local strategy for email/password, jwt strategy for access token verification, and oauth strategies for external providers. It is important not to mix concepts: OAuth is a delegated authorization protocol, JWT is a token format, and Passport is middleware/framework support for authentication strategies.",
            ("bullets", [
                "OAuth: external provider or delegated access.",
                "JWT: internal access token for the application's own API.",
                "Refresh tokens should be stored more safely than JS-accessible localStorage.",
                "Authorization checks must still happen in the backend/service layer."
            ]),
        ],
    )

    add_section(doc, "DevOps / AWS / Docker")

    add_qa(
        doc,
        23,
        "Docker image для backend стал весить 1.5 GB. Как будешь уменьшать?\nПроверяет: Docker multi-stage,\u00a0.dockerignore, prod deps.",
        "A Docker image for the backend became 1.5 GB. How would you reduce it?\nChecks: Docker multi-stage, .dockerignore, production dependencies.",
        [
            "Сначала нужно понять, что именно попало в image: node_modules, исходники, тесты, build cache, .git, временные файлы, локальные uploads, sourcemaps, dev dependencies. Для этого можно использовать docker history, dive или анализ слоев. Без анализа легко уменьшить не тот слой.",
            "Для Node/Nest backend стандартное решение — multi-stage build. На build stage устанавливаются все зависимости и компилируется TypeScript. В runtime stage копируются только compiled dist, package files и production dependencies. Dev dependencies, tests, исходные TS-файлы и build artifacts не должны попадать в финальный image, если они не нужны для запуска.",
            ".dockerignore обязателен: нужно исключить node_modules, .git, coverage, logs, local env files, test fixtures, uploads и другие лишние директории. Также полезно использовать slim base image и правильно упорядочить COPY, чтобы кешировать npm install. Но alpine/slim нужно выбирать осознанно: некоторые native modules могут требовать дополнительных системных библиотек.",
            ("bullets", [
                "Проверить слои через docker history/dive.",
                "Использовать multi-stage build.",
                "В runtime image оставить только dist и production dependencies.",
                "Настроить .dockerignore.",
                "Не копировать тесты, кэш, .git и локальные файлы.",
                "Проверить, что image не содержит secrets."
            ]),
        ],
        [
            "First, identify what actually got into the image: node_modules, source files, tests, build cache, .git, temporary files, local uploads, sourcemaps, or dev dependencies. Tools like docker history, dive, or layer analysis can help. Without analysis, it is easy to optimize the wrong layer.",
            "For a Node/Nest backend, the standard solution is a multi-stage build. The build stage installs all dependencies and compiles TypeScript. The runtime stage copies only compiled dist, package files, and production dependencies. Dev dependencies, tests, source TS files, and build artifacts should not be present in the final image unless they are needed at runtime.",
            ".dockerignore is mandatory: exclude node_modules, .git, coverage, logs, local env files, test fixtures, uploads, and other unnecessary directories. A slim base image and a good COPY order can also help cache npm install. However, alpine/slim should be chosen deliberately because some native modules require additional system libraries.",
            ("bullets", [
                "Inspect layers with docker history/dive.",
                "Use a multi-stage build.",
                "Keep only dist and production dependencies in the runtime image.",
                "Configure .dockerignore.",
                "Do not copy tests, cache, .git, or local files.",
                "Verify that the image does not contain secrets."
            ]),
        ],
    )

    add_qa(
        doc,
        24,
        "Какие IAM permissions нужны сервису, который читает/пишет media в S3, но не должен иметь полный доступ к bucket?\nПроверяет: IAM least privilege.",
        "What IAM permissions are needed for a service that reads/writes media in S3 but must not have full access to the bucket?\nChecks: IAM least privilege.",
        [
            "Нужно применить принцип least privilege: сервис получает только те действия и только на те ресурсы, которые ему реально нужны. Если сервис работает с media конкретного приложения или tenant-а, policy лучше ограничить prefix-ом, например arn:aws:s3:::bucket/media/app-name/* или tenant-specific prefix.",
            "Для чтения и записи обычно нужны s3:GetObject и s3:PutObject на object resources. Если сервис удаляет или заменяет файлы, может понадобиться s3:DeleteObject. Для multipart upload могут понадобиться дополнительные permissions, связанные с multipart lifecycle. Для listing иногда нужен s3:ListBucket, но его нужно ограничивать условием s3:prefix, а не давать весь bucket.",
            "Не стоит выдавать s3:* или доступ ко всем buckets. Также важно настроить bucket policy, encryption, ownership, block public access и не хранить AWS keys в image или коде. В ECS лучше использовать task role, чтобы сервис получал временные credentials через IAM role, а не через статичные secrets.",
            ("bullets", [
                "Разрешить только нужные actions: GetObject, PutObject, возможно DeleteObject.",
                "Ограничить resources конкретным bucket/prefix.",
                "ListBucket давать только при необходимости и с prefix condition.",
                "Использовать ECS task role вместо статичных AWS keys.",
                "Не давать s3:* без крайней необходимости."
            ]),
        ],
        [
            "The least privilege principle should be applied: the service receives only the actions and resources it actually needs. If the service works with media for a specific application or tenant, the policy should be restricted by prefix, for example arn:aws:s3:::bucket/media/app-name/* or a tenant-specific prefix.",
            "For reading and writing, s3:GetObject and s3:PutObject are usually required on object resources. If the service deletes or replaces files, s3:DeleteObject may be needed. Multipart upload may require additional permissions related to multipart lifecycle. Listing sometimes requires s3:ListBucket, but it should be restricted with an s3:prefix condition, not granted for the whole bucket.",
            "The service should not receive s3:* or access to all buckets. Bucket policy, encryption, ownership, and block public access should also be configured. AWS keys should not be stored in the image or code. In ECS, a task role is preferable so the service receives temporary credentials through IAM rather than static secrets.",
            ("bullets", [
                "Allow only required actions: GetObject, PutObject, possibly DeleteObject.",
                "Restrict resources to a specific bucket/prefix.",
                "Grant ListBucket only when needed and with a prefix condition.",
                "Use an ECS task role instead of static AWS keys.",
                "Do not grant s3:* unless there is a very strong reason."
            ]),
        ],
    )

    add_qa(
        doc,
        25,
        "Какую CloudWatch/monitoring картину ты хочешь видеть для AI processing системы?\nПроверяет: metrics, alerts, observability.",
        "What CloudWatch/monitoring picture would you want to see for an AI processing system?\nChecks: metrics, alerts, observability.",
        [
            "Для AI processing системы важны не только CPU и memory. Нужна картина всего pipeline: сколько jobs создано, сколько в очереди, сколько выполняется, сколько завершилось успешно, сколько упало, сколько попало в DLQ. Отдельно нужно видеть duration по этапам: upload, transcription, summarization, segmentation, saving results, notification.",
            "Технические метрики: API latency/error rate, Kafka consumer lag, retry count, DLQ size, PostgreSQL pool usage, slow queries, Redis latency, S3 errors, ECS task health, memory, CPU, event loop delay. Для WebSocket — количество активных соединений, reconnect rate, ошибки доставки событий.",
            "AI-специфичные метрики: tokens in/out, cost per job/user/org, OpenAI latency, model error rate, invalid JSON rate, fallback/retry rate, prompt version distribution. Нужны alerts не только на падение сервиса, но и на бизнесовые аномалии: рост стоимости, всплеск failed jobs, очередь не разгребается, p95 duration выше нормы.",
            ("bullets", [
                "Dashboard должен показывать здоровье pipeline, а не только серверов.",
                "Нужны correlationId/jobId в логах.",
                "Метрики очередей и DLQ критичны для async-системы.",
                "AI cost и token usage должны быть first-class метриками.",
                "Alerts должны быть actionable, а не просто шумом."
            ]),
        ],
        [
            "For an AI processing system, CPU and memory are not enough. The dashboard should show the whole pipeline: how many jobs were created, queued, running, completed successfully, failed, and sent to DLQ. Stage duration should be visible separately: upload, transcription, summarization, segmentation, result saving, and notification.",
            "Technical metrics include API latency/error rate, Kafka consumer lag, retry count, DLQ size, PostgreSQL pool usage, slow queries, Redis latency, S3 errors, ECS task health, memory, CPU, and event loop delay. For WebSockets, monitor active connections, reconnect rate, and event delivery errors.",
            "AI-specific metrics include tokens in/out, cost per job/user/org, OpenAI latency, model error rate, invalid JSON rate, fallback/retry rate, and prompt version distribution. Alerts should cover not only service crashes, but also business anomalies: cost spikes, failed job spikes, queue backlog growth, or p95 job duration exceeding normal thresholds.",
            ("bullets", [
                "The dashboard should show pipeline health, not only server health.",
                "Logs need correlationId/jobId.",
                "Queue and DLQ metrics are critical for async systems.",
                "AI cost and token usage should be first-class metrics.",
                "Alerts should be actionable, not just noisy."
            ]),
        ],
    )

    doc.save(OUT_PATH)


if __name__ == "__main__":
    build()
